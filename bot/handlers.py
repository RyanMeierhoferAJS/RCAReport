import io
import logging
from telegram import Update
from telegram.ext import ContextTypes
from ai.router import extract_from_message, extract_pdp_from_document
from modules.search import search_and_answer
from db import client as db
from bot.utils import safe_reply as _safe_reply

logger = logging.getLogger(__name__)

# How many recent turns to send to Claude (full history stored in Supabase)
_HISTORY_WINDOW = 20


def _get_history(chat_id: int) -> list[dict]:
    try:
        return db.get_conversation_history(chat_id, limit=_HISTORY_WINDOW)
    except Exception:
        logger.warning("Failed to load conversation history for chat %s", chat_id)
        return []


def _save_history(chat_id: int, user_text: str, assistant_text: str) -> None:
    try:
        db.append_conversation(chat_id, "user", user_text)
        db.append_conversation(chat_id, "assistant", assistant_text)
    except Exception:
        logger.warning("Failed to save conversation history for chat %s", chat_id)


_QUESTION_STARTERS = (
    "what", "who", "when", "where", "how", "which", "why",
    "show", "list", "find", "tell", "have i", "did i",
    "do i have", "am i", "any ", "meetings", "calendar",
    "tasks", "decisions", "ideas", "remind", "summarise", "summarize",
)


def _looks_like_question(text: str) -> bool:
    lower = text.lower().strip()
    return lower.endswith("?") or lower.startswith(_QUESTION_STARTERS)


async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    text       = update.message.text
    message_id = update.message.message_id
    chat_id    = update.message.chat_id

    capture_id = db.store_capture(
        raw_text=text,
        telegram_msg_id=message_id,
        telegram_chat_id=chat_id,
        media_type="text",
    )

    history = _get_history(chat_id)

    # Fast-path for obvious questions — skip Tier 1 extraction
    if _looks_like_question(text):
        result = search_and_answer(text, history)
        _save_history(chat_id, text, result)
        await _safe_reply(update, result)
        return

    # Tier 1: classify and extract structured entities
    extracted = extract_from_message(text)

    # If AI also classified it as a question, answer it
    if extracted.get("classification") == "question":
        result = search_and_answer(text, history)
        _save_history(chat_id, text, result)
        await _safe_reply(update, result)
        return

    stored = []

    for task in extracted.get("tasks", []):
        db.create_task(
            title=task["title"],
            description=task.get("description"),
            priority=task.get("priority", "medium"),
            due_date=task.get("due_date"),
            project=task.get("project"),
            source_capture_id=capture_id,
        )
        stored.append(f"✅ *Task:* {task['title']}")

    for decision in extracted.get("decisions", []):
        db.create_decision(
            title=decision["title"],
            description=decision.get("description"),
            reason=decision.get("reason"),
            alternatives=decision.get("alternatives", []),
            project=decision.get("project"),
            source_capture_id=capture_id,
        )
        stored.append(f"🎯 *Decision:* {decision['title']}")

    for achievement in extracted.get("achievements", []):
        db.create_career_event(
            type=achievement.get("type", "achievement"),
            title=achievement["title"],
            description=achievement.get("description"),
            value_pounds=achievement.get("value_pounds"),
            project=achievement.get("project"),
            source_capture_id=capture_id,
        )
        stored.append(f"🏆 *Achievement:* {achievement['title']}")

    for note in extracted.get("notes", []):
        db.create_note(
            content=note["content"],
            tags=note.get("tags", []),
            entities=note.get("entities", []),
            project=note.get("project"),
            source_capture_id=capture_id,
        )

    for idea in extracted.get("ideas", []):
        db.create_idea(
            title=idea["title"],
            description=idea.get("description"),
            category=idea.get("category", "general"),
            project=idea.get("project"),
            source_capture_id=capture_id,
        )
        stored.append(f"💡 *Idea:* {idea['title']}")

    for completion in extracted.get("completions", []):
        partial = completion.get("title", "").strip()
        if not partial:
            continue
        matches = db.find_open_tasks_by_title(partial)
        if not matches:
            open_tasks = db.get_open_tasks()
            lower = partial.lower()
            matches = [t for t in open_tasks if lower in t["title"].lower()]
        if len(matches) == 1:
            db.complete_task_by_id(matches[0]["id"])
            stored.append(f"✅ *Completed:* {matches[0]['title']}")
        elif len(matches) > 1:
            titles = ", ".join(t["title"] for t in matches[:3])
            stored.append(f"⚠️ Multiple tasks match '{partial}' — use /done to pick one ({titles}…)")

    for pdp_ev in extracted.get("pdp_evidence", []):
        action_title = pdp_ev.get("action_title", "")
        evidence_text = pdp_ev.get("evidence", "")
        exceeded = pdp_ev.get("exceeded", False)
        if action_title and evidence_text:
            actions = db.get_pdp_actions()
            matched = next(
                (a for a in actions if action_title.lower() in a["title"].lower()),
                None,
            )
            if matched:
                new_status = "exceeded" if exceeded else "in_progress"
                db.add_pdp_evidence(matched["id"], evidence_text, new_status)
                stored.append(f"📈 *PDP evidence logged* for: {matched['title']}")

    # Update last_activity on any mentioned projects
    for project in extracted.get("projects_mentioned", []):
        try:
            db.update_project_activity(project)
        except Exception:
            pass

    # Reply: use AI's natural response if available, else summarise what was stored
    reply = extracted.get("response", "")
    if not reply and stored:
        reply = "\n".join(stored)
    elif not reply:
        reply = "Got it, stored."

    _save_history(chat_id, text, reply)
    await _safe_reply(update, reply)


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    doc = update.message.document
    fname = doc.file_name or ""

    if not (fname.endswith(".docx") or fname.endswith(".doc")):
        await update.message.reply_text(
            "Send me a `.docx` file and I'll extract your PDP actions from it automatically.\n"
            "_(PDF support coming soon)_",
            parse_mode="Markdown",
        )
        return

    await update.message.reply_text("_Reading your PDP document…_", parse_mode="Markdown")

    try:
        import docx  # python-docx

        tg_file = await doc.get_file()
        buf = io.BytesIO()
        await tg_file.download_to_memory(buf)
        buf.seek(0)

        document = docx.Document(buf)

        # Extract paragraphs
        para_text = "\n".join(p.text for p in document.paragraphs if p.text.strip())

        # Extract tables (PDPs are almost always in table format)
        table_lines = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                # Deduplicate merged cells (python-docx repeats merged cell text)
                seen = []
                for c in cells:
                    if not seen or c != seen[-1]:
                        seen.append(c)
                if seen:
                    table_lines.append(" | ".join(seen))

        table_text = "\n".join(table_lines)
        full_text = "\n\n".join(filter(None, [para_text, table_text]))

        if not full_text.strip():
            await update.message.reply_text("The document appears to be empty or unreadable.")
            return

        logger.info("PDP document extracted %d chars (%d para, %d table rows)",
                    len(full_text), len(para_text.splitlines()), len(table_lines))

        await update.message.reply_text("_Extracting PDP actions with AI…_", parse_mode="Markdown")
        result = extract_pdp_from_document(full_text)

        actions = result.get("pdp_actions", [])
        if not actions:
            preview = full_text[:500].replace("_", "\\_").replace("*", "\\*")
            await update.message.reply_text(
                "Couldn't extract PDP actions — here's what I read from the document "
                "(first 500 chars):\n\n"
                f"`{preview}`\n\n"
                "If this looks wrong, the document may use a format I can't read yet. "
                "Try copy-pasting the content as a plain text message instead.",
                parse_mode="Markdown",
            )
            return

        saved = []
        for action in actions:
            if not action.get("title"):
                continue
            db.create_pdp_action(
                title=action["title"],
                description=action.get("description"),
                category=action.get("category", "general"),
                objective=action.get("objective"),
                target_date=action.get("target_date"),
            )
            cat = action.get("category", "general").capitalize()
            saved.append(f"✅ [{cat}] {action['title']}")

        summary = result.get("summary", "")
        header = f"*PDP imported — {len(saved)} actions loaded*"
        if summary:
            header += f"\n_{summary}_"

        reply = header + "\n\n" + "\n".join(saved)
        reply += "\n\nUse `/pdp` to view your dashboard, or `/pdp analyse` for an AI gap analysis."
        await update.message.reply_text(reply, parse_mode="Markdown")

    except Exception:
        logger.exception("Document PDP import failed")
        await update.message.reply_text("Something went wrong reading the document — please try again.")


async def handle_voice(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "Voice notes are coming soon — send as text for now."
    )


async def error_handler(update: object, context: ContextTypes.DEFAULT_TYPE) -> None:
    err = context.error
    logger.error("Unhandled exception", exc_info=err)
    if isinstance(update, Update) and update.message:
        err_type = type(err).__name__
        err_msg  = str(err)[:200]
        await update.message.reply_text(
            f"Error ({err_type}): {err_msg}\n\nCheck Railway logs for full traceback."
        )
